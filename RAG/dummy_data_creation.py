import os
import pandas as pd
from docx import Document

# 1. Initialize and build the directory tree
folders = ["behavioral_anomalies", "network_typologies", "legal_compliance"]
for folder in folders:
    os.makedirs(folder, exist_ok=True)

print("📁 Folders verified and initialized.")

# ----------------------------------------------------
# 1. BEHAVIORAL ANOMALIES MOCK FILES
# ----------------------------------------------------
# File A: Plain Text summary of bot/ATO behavior
with open("behavioral_anomalies/owasp_ato_guide.txt", "w", encoding="utf-8") as f:
    f.write(
        "OWASP Automated Threats Section OAT-019: Account Takeover (ATO).\n"
        "Indicators include extreme typing speed anomalies, immediate transaction execution "
        "following an unauthorized login, and session durations dropping below 60 seconds. "
        "Credential stuffing (OAT-008) often precedes this exploit, characterized by automated bots "
        "testing large variations of usernames and passwords against banking API endpoints."
    )

# File B: Tabular CSV representing anomalous user telemetry logs
behavior_data = {
    "session_id": ["SESS_9921", "SESS_1044", "SESS_8831"],
    "typing_speed_flag": [True, False, True],
    "session_duration_minutes": [0.4, 12.5, 0.8],
    "failed_transaction_count_24h": [9, 0, 14],
    "account_takeover_suspected": [True, False, True]
}
pd.DataFrame(behavior_data).to_csv("behavioral_anomalies/session_telemetry_sample.csv", index=False)


# ----------------------------------------------------
# 2. NETWORK TYPOLOGIES MOCK FILES
# ----------------------------------------------------
# File A: Word Document detailing money mule network frameworks
doc = Document()
doc.add_heading("Europol Money Muling Framework & Financial Ring Typologies", level=1)
doc.add_paragraph(
    "Money mule rings operate as decentralized networks where multiple proxy identities "
    "route illicit funds through intermediate accounts. A classic structural graph signature "
    "is a high many-to-one mapping layout where distinct, seemingly unlinked customer nodes "
    "share a single hardware device fingerprint (mule_device) or routing beneficiary parameters (mule_bene)."
)
doc.save("network_typologies/europol_mule_report.docx")

# File B: Excel spreadsheet mapping structural entities inside a fraud cluster
network_data = {
    "node_id": ["CUST_45A", "CUST_92B", "DEV_MULE_BASE"],
    "entity_type": ["Customer Persona", "Customer Persona", "Shared Blacklisted Device"],
    "relationship_risk_score": [95.50, 91.20, 100.00],
    "cluster_assignment": ["Syndicate_Alpha", "Syndicate_Alpha", "Syndicate_Alpha"]
}
pd.DataFrame(network_data).to_excel("network_typologies/graph_cluster_mapping.xlsx", index=False)


# ----------------------------------------------------
# 3. LEGAL COMPLIANCE MOCK FILES
# ----------------------------------------------------
# File A: Plain Text tracking FinCEN regulatory guidelines
with open("legal_compliance/fincen_structuring_red_flags.txt", "w", encoding="utf-8") as f:
    f.write(
        "FinCEN Suspicious Activity Quick Reference Guide - Structuring Guidelines.\n"
        "Individuals or syndicates attempting to evade standard cash transaction reporting thresholds "
        "often split large volumes of funds into smaller, sequential increments (smurfing behavior). "
        "For example, processing three separate $4,500 wire transfers within an hour instead of a single "
        "$13,500 transfer triggers direct structuring alerts. Transactions routing to high-risk "
        "sanctioned regions such as North Korea or Russia violate international policy and mandate an immediate SAR."
    )

# File B: CSV tracking localized high-risk geographic flags
compliance_data = {
    "country_code": ["KP", "RU", "IR", "SY"],
    "country_name": ["North Korea", "Russia", "Iran", "Syria"],
    "sanction_status": ["Critical", "Critical", "High", "High"],
    "regulatory_reference": ["OFAC-2026-KP", "EU-2026-RU", "UN-RE-092", "UKHM-991"]
}
pd.DataFrame(compliance_data).to_csv("legal_compliance/sanction_lookup_matrix.csv", index=False)

print("✨ Mock test documents successfully written across all folders!")